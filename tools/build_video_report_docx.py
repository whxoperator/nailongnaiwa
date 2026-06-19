from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "奶龙奶蛙图像分类项目报告.docx"

IMAGES = [
    (
        ROOT / "nailong_naiwa_splits" / "test" / "nailong" / "021697702C6C435DB551A735B36A02ED_17.jpg",
        "图 1  奶龙样例：形象特征相对清晰，适合用于模型学习典型外观。",
    ),
    (
        ROOT / "nailong_naiwa_splits" / "test" / "naiwa" / "naiwa_from_nailong_zip_001.jpg",
        "图 2  奶蛙样例：场景复杂、风格接近真实网络图片，适合检验泛化能力。",
    ),
    (
        ROOT / "nailong_naiwa_splits" / "test" / "naiwa" / "625e2cd7f56c3f5e991d0bed77ceb8d9bd5f659f.jpg",
        "图 3  奶蛙样例：带有文字与强背景效果，体现数据来源的多样性。",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for run_item in paragraph.runs:
        run_item.font.name = "Calibri"
        run_item._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run_item.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(46, 116, 181 if level == 1 else 120)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix) :]
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("奶龙 / 奶蛙图像分类项目报告")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("结合项目讲稿、功能演示与样例图片整理")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "一、项目背景与必要性", 1)
    add_para(
        doc,
        "在今天的互联网环境里，图片已经不只是普通的娱乐内容，它本质上是一种信息载体。表情包、角色图片、截图和二创内容背后，都涉及同一个基础问题：机器能不能理解图像，能不能把视觉信息转化成明确、可判断、可管理的数据。",
    )
    add_para(
        doc,
        "图像分类是这个问题中非常基础、也非常必要的一步。只有先完成“这是什么”的判断，后面才谈得上搜索、推荐、归档、数据清洗、内容管理以及更复杂的智能应用。因此，本项目虽然以奶龙和奶蛙为对象，但背后对应的是一个更宏观的任务：让计算机从大量相似图片中提取差异，并给出稳定判断。",
    )
    add_para(
        doc,
        "具体到本任务，很多人在真实网络图片中分不清奶龙和奶蛙。二者形象风格接近，传播来源复杂，图片质量也不统一。如果完全依靠人工逐张判断，不仅效率低，也容易产生误差。因此，构建一个自动识别工具不是可有可无的装饰，而是后续数据整理、功能扩展和智能分析的必要基础。",
    )

    add_heading(doc, "二、数据来源与样例展示", 1)
    add_para(
        doc,
        "本项目的数据主要来自两个部分：一部分参考并使用 GitHub 开源项目 spawner1145/NailongRecognize，另一部分来自小红书等平台下载并整理得到的图片。由于图片来源不完全一致，项目中专门划分了 train、test 和 generalization 三类数据集，用来观察模型在不同来源和风格图片上的表现。",
    )

    for image_path, caption in IMAGES:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Cm(8.2))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        for run in cap.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "三、项目功能设计", 1)
    add_para(
        doc,
        "项目并不是只训练一个模型，而是包含数据预处理、CNN 训练、预测界面、传统图像算法对比和评估报告输出的一整套流程。它更像是一个缩小版的图像识别系统，可以把抽象的算法概念落到可操作、可验证的实际任务中。",
    )
    add_bullet(doc, "训练模型：使用轻量级 CNN，对奶龙和奶蛙图片进行二分类训练，并将模型保存到 models 目录。")
    add_bullet(doc, "数据增强：通过随机裁剪、翻转、旋转和颜色扰动，减少模型对固定图片样式的依赖。")
    add_bullet(doc, "水印处理：对部分图片进行水印区域预处理，尽量避免模型只记住角落或水印特征。")
    add_bullet(doc, "本地 Web UI：支持图片上传、模型选择、算法选择、预测结果展示和置信度提示。")
    add_bullet(doc, "算法对比：支持 CNN、颜色均值、RGB 直方图、缩略图 kNN、边缘直方图等方法的横向比较。")

    add_heading(doc, "四、功能演示规划", 1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["时间段", "建议时长", "画面内容", "讲解重点"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    rows = [
        ("0:00-0:55", "55 秒", "展示项目目录、README、GitHub 页面", "从图像信息理解讲到项目必要性"),
        ("0:55-1:45", "50 秒", "展示数据集目录和数据分布", "说明数据来源与划分方式"),
        ("1:45-2:25", "40 秒", "展示训练脚本和模型文件", "说明 CNN 模型训练流程"),
        ("2:25-3:35", "70 秒", "打开 Web UI 并上传图片", "展示核心预测功能"),
        ("3:35-4:15", "40 秒", "展示 Compare all algorithms 和报告", "展示多算法对比结果"),
        ("4:15-5:05", "50 秒", "展示 README、GitHub 或项目目录", "说明不足和开发困难"),
        ("5:05-5:40", "35 秒", "回到项目整体结构或 UI 页面", "总结项目价值和未来方向"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    add_heading(doc, "五、模型效果与算法对比", 1)
    add_para(
        doc,
        "根据当前测试集报告，平衡数据训练得到的 CNN 模型表现最好，测试集准确率约为 90.24%；缩略图 kNN 准确率约为 87.80%；RGB 直方图原型准确率约为 85.37%。这说明 CNN 模型整体效果较好，但传统图像方法在小数据集场景下也具备一定参考价值。",
    )
    add_para(
        doc,
        "多算法对比功能的意义在于，它可以让同一张图片接受不同算法的判断。如果多个算法都指向同一类别，结果通常更有参考价值；如果算法之间分歧明显，则说明该图片可能存在风格特殊、画质较低或样本特征不稳定等问题。",
    )

    add_heading(doc, "六、项目不足", 1)
    add_bullet(doc, "数据集规模较小，测试集数量有限，准确率存在一定偶然性。")
    add_bullet(doc, "分类类别目前只有奶龙和奶蛙，无法处理其他角色或无关图片。")
    add_bullet(doc, "模型可能受到背景、水印、图片压缩质量和图片来源的影响，泛化能力仍需更多真实图片验证。")
    add_bullet(doc, "Web UI 目前更偏实验演示，还缺少批量上传、历史记录、用户管理和人工审核等完整功能。")
    add_bullet(doc, "由于时间原因，原本计划中的云端数据库、多人协作数据管理和预测历史保存功能没有完全实现。")
    add_bullet(doc, "这是第一次比较完整地上传相关项目到 GitHub，仓库整理、README 编写和提交流程也消耗了较多时间。")

    add_heading(doc, "七、后续改进方向", 1)
    add_bullet(doc, "继续扩充数据集，增加不同来源、不同画质和不同背景的图片。")
    add_bullet(doc, "加入错误样本分析，对模型容易混淆的图片进行针对性补充训练。")
    add_bullet(doc, "优化 Web UI，加入批量预测、历史记录、人工修正标签和自动生成评估报告功能。")
    add_bullet(doc, "建立云端数据库，保存用户上传图片、预测结果、人工修正标签和历史记录。")
    add_bullet(doc, "引入更多评估指标，例如精确率、召回率、F1 分数和泛化集专项报告。")

    add_heading(doc, "八、总结", 1)
    add_para(
        doc,
        "总体来看，本项目完成了从数据准备、模型训练、单图预测、算法对比到报告输出的一整套流程。它虽然规模不大，但功能闭环较完整，能够围绕“奶龙还是奶蛙”这个具体问题完成一次真实的图像分类实践。",
    )
    add_para(
        doc,
        "对本次任务而言，项目的价值不仅在于得到一个可以运行的分类器，也在于完整经历了数据收集、模型训练、功能展示、效果评估和 GitHub 上传的过程。虽然仍有不少功能可以继续扩展，但它已经为后续构建更完整的在线识别系统打下了基础。",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Nailong / Naiwa Image Classifier 项目报告")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
