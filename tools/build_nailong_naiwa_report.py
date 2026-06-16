from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("nailong_naiwa_report.docx")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(102, 112, 133)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("奶龙 / 奶蛙识别实验报告")
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("奶龙 / 奶蛙图像识别神经网络实验报告")
    set_run_font(run, size=22, bold=True, color=RGBColor(11, 37, 69))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("数据预处理、平衡采样、训练增强与交互式识别系统")
    set_run_font(run, size=12, color=MUTED)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="E5E7EB")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=True, color=RGBColor(31, 58, 95))
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_fixed_table_width(table, widths)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
        for cell in cells:
            set_cell_margins(cell)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(row_idx == 0))
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_table(doc, ["项目", "内容"], [[key, value] for key, value in rows], [2.05, 4.35])


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(item)
        set_run_font(run)


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(item)
        set_run_font(run)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_callout(
        doc,
        "本实验完成了从数据整理、去水印弱化、平衡采样、神经网络训练到网页交互识别的完整流程。"
        "最终生成的新模型 best test accuracy 为 0.902。",
    )

    add_kv_table(
        doc,
        [
            ("项目路径", r"D:\nndl\nailong"),
            ("任务目标", "识别图片属于 nailong（奶龙）还是 naiwa（奶蛙）"),
            ("最终模型", "models/nailong_naiwa_balanced_cnn.pt"),
            ("交互界面", "http://127.0.0.1:8770"),
            ("报告日期", "2026-06-16"),
        ],
    )

    doc.add_heading("1. 实验背景与目标", level=1)
    doc.add_paragraph(
        "本实验面向奶龙与奶蛙两类图像的二分类识别任务。原始数据中，奶龙图片数量明显多于奶蛙图片，"
        "且奶蛙图片右下角普遍存在平台水印。若直接训练，模型容易学习到类别无关的水印特征，而不是主体外观特征。"
    )
    doc.add_paragraph(
        "因此，本次工作重点不是单纯训练一个模型，而是建立一套更稳健的数据处理与反馈闭环：先弱化水印影响，"
        "再通过平衡采样和数据增强训练模型，最后提供可选择模型的交互界面，并支持将人工确认正确的新图片继续纳入数据集。"
    )

    doc.add_heading("2. 数据整理与预处理", level=1)
    add_table(
        doc,
        ["数据项", "数量 / 路径", "说明"],
        [
            ["原始 nailong 图片", "345 张", "来自 nailong_naiwa_10_demo/nailong"],
            ["实验 nailong 图片", "100 张", "从 345 张中按随机种子抽样，避免类别不平衡"],
            ["原始 naiwa 图片", "105 张", "来自 nailong_naiwa_10_demo/naiwa"],
            ["预处理 naiwa 图片", "105 张", "输出到 nailong_naiwa_10_demo/naiwa_preprocessed"],
        ],
        [1.55, 1.75, 3.1],
    )

    doc.add_heading("2.1 水印弱化处理", level=2)
    doc.add_paragraph(
        "奶蛙图片右下角存在较多水印。预处理脚本 preprocess_naiwa_watermarks.py 对右下区域建立遮罩，"
        "并使用局部平滑与模糊填补方式弱化水印。该处理保留原图不变，另存为 naiwa_preprocessed，用于后续训练。"
    )
    add_bullets(
        doc,
        [
            "输入目录：nailong_naiwa_10_demo/naiwa",
            "输出目录：nailong_naiwa_10_demo/naiwa_preprocessed",
            "预览文件：nailong_naiwa_10_demo/naiwa_watermark_preview.jpg",
            "目的：减少模型把平台水印等角落特征误学为 naiwa 类别特征的风险。",
        ],
    )

    doc.add_heading("2.2 平衡采样与数据拆分", level=2)
    doc.add_paragraph(
        "由于 nailong 图片远多于 naiwa 图片，本次实验按要求只从 nailong 中随机抽取 100 张，"
        "与 105 张 naiwa 预处理图片组成平衡实验集。随后将两类图片分别随机拆分为训练集、测试集和泛化集。"
    )
    add_table(
        doc,
        ["集合", "nailong", "naiwa", "用途"],
        [
            ["train", "60", "63", "模型训练"],
            ["test", "20", "21", "训练过程中的测试评估"],
            ["generalization", "20", "21", "额外泛化观察"],
            ["合计", "100", "105", "平衡实验集"],
        ],
        [1.3, 1.1, 1.1, 2.9],
    )

    doc.add_heading("3. 模型与训练方法", level=1)
    doc.add_paragraph(
        "模型使用 PyTorch 实现的小型卷积神经网络 SmallImageCNN。该模型由多层卷积、BatchNorm、ReLU、"
        "池化和全局平均池化组成，最后使用 Dropout 与线性分类层输出两个类别概率。"
    )
    add_table(
        doc,
        ["方法", "实现位置", "作用"],
        [
            ["方法 3：通用数据增强", "nailong_model.py / ClassAwareImageFolder", "随机裁剪、翻转、旋转、颜色扰动，提高小数据集泛化能力"],
            ["方法 2：右下角随机遮挡/模糊", "nailong_model.py / RandomCornerOcclusion", "对 nailong 训练图也扰动右下角，降低模型依赖角落水印特征的可能性"],
            ["类别平衡", "train_nailong_naiwa.py", "只使用 100 张 nailong，与 105 张 naiwa 保持接近"],
            ["标签平滑", "train_nailong_naiwa.py", "CrossEntropyLoss(label_smoothing=0.04)，降低过拟合"],
        ],
        [1.65, 2.25, 2.5],
    )
    add_kv_table(
        doc,
        [
            ("训练命令", "python train_nailong_naiwa.py --cpu --epochs 12 --repeats 3 --batch-size 16 --image-size 128 --print-every 3 --out models/nailong_naiwa_balanced_cnn.pt"),
            ("设备", "CPU"),
            ("训练轮数", "12 epochs"),
            ("批大小", "16"),
            ("图像尺寸", "128 x 128"),
            ("输出模型", "models/nailong_naiwa_balanced_cnn.pt"),
        ],
    )

    doc.add_heading("4. 实验结果", level=1)
    doc.add_paragraph("本次训练在 CPU 上完成，训练脚本会保存测试集准确率最高的模型状态。快速实验的最佳测试集准确率为 0.902。")
    add_table(
        doc,
        ["Epoch", "Loss", "Train Acc", "Test Acc", "Gen Acc"],
        [
            ["1 / 12", "0.5991", "0.659", "0.780", "0.732"],
            ["3 / 12", "0.4905", "0.770", "0.854", "0.951"],
            ["6 / 12", "0.4849", "0.791", "0.902", "0.927"],
            ["9 / 12", "0.4233", "0.856", "0.878", "0.927"],
            ["12 / 12", "0.4137", "0.854", "0.854", "0.902"],
        ],
        [1.2, 1.1, 1.3, 1.3, 1.3],
    )
    add_callout(
        doc,
        "结论：在平衡采样和增强策略下，模型已经可以作为奶龙 / 奶蛙识别实验原型使用。"
        "由于数据量仍有限，后续继续补充无水印、多角度、多背景图片会进一步提升可靠性。",
    )

    doc.add_heading("5. 交互式识别界面", level=1)
    doc.add_paragraph(
        "交互界面由 predict_nailong_naiwa_ui.py 提供，启动后会读取 models 目录下的 .pt 模型文件。"
        "用户可以选择不同模型，上传新图片，查看预测类别与概率。如果用户确认判断正确，界面会把该图片随机复制到"
        " train、test 或 generalization 对应类别目录。"
    )
    add_steps(
        doc,
        [
            "运行命令：python predict_nailong_naiwa_ui.py --host 127.0.0.1 --port 8770 --cpu",
            "在浏览器打开 http://127.0.0.1:8770。",
            "选择模型并上传新图片。",
            "查看 nailong（奶龙）与 naiwa（奶蛙）的概率结果。",
            "如果判断正确，点击“判断正确，随机加入数据集”。",
        ],
    )
    add_table(
        doc,
        ["功能", "说明"],
        [
            ["多模型选择", "兼容当前新模型和旧模型，可从 models/*.pt 选择"],
            ["预测展示", "显示预测类别、中文类别名和两个类别概率"],
            ["人工确认入库", "确认正确后复制图片，不删除原上传文件"],
            ["随机归类", "按权重随机进入 train / test / generalization"],
        ],
        [1.9, 4.5],
    )

    doc.add_heading("6. 覆盖与新增的主要文件", level=1)
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["preprocess_naiwa_watermarks.py", "批量弱化 naiwa 图片右下角水印，输出 naiwa_preprocessed"],
            ["nailong_model.py", "模型结构、图像增强、预测变换、模型保存与加载"],
            ["train_nailong_naiwa.py", "平衡采样、数据拆分、训练与评估"],
            ["predict_nailong_naiwa_ui.py", "交互式网页识别与人工确认入库"],
            ["README_nailong_naiwa_classifier.md", "中文使用说明与实验记录"],
        ],
        [2.55, 3.85],
    )

    doc.add_heading("7. 风险与后续改进建议", level=1)
    add_bullets(
        doc,
        [
            "数据量仍然偏小，建议继续补充更多无水印、不同角度、不同背景的奶蛙图片。",
            "当前去水印是训练友好的弱化处理，不等同于精修级图像修复。",
            "如果后续加入更多图片，应重新运行训练脚本，生成新的模型版本并在交互界面中比较。",
            "建议额外保留一批完全未参与训练和调参的真实新图片，作为最终人工验收集。",
            "若有 GPU，可提高 epochs 和 repeats，训练更充分的模型版本。",
        ],
    )

    doc.add_heading("8. 总结", level=1)
    doc.add_paragraph(
        "本次工作已经完成可运行的奶龙 / 奶蛙图像识别原型：数据经过水印弱化和平衡采样，训练阶段加入了"
        "通用增强与右下角扰动，模型达到 0.902 的最佳测试集准确率，并提供了可选择模型、可反馈入库的交互式网页界面。"
        "该系统已经具备继续迭代数据和模型的基础闭环。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
